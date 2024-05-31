import os
import argparse
import io
from typing import List
import time

import pypdfium2
import streamlit as st
from surya.detection import batch_text_detection
from surya.layout import batch_layout_detection
from surya.model.detection.segformer import load_model, load_processor
from surya.model.recognition.model import load_model as load_rec_model
from surya.model.recognition.processor import load_processor as load_rec_processor
from surya.model.ordering.processor import load_processor as load_order_processor
from surya.model.ordering.model import load_model as load_order_model
from surya.ordering import batch_ordering
from surya.postprocessing.heatmap import draw_polys_on_image
from surya.ocr import run_ocr
from surya.postprocessing.text import draw_text_on_image
from PIL import Image
from surya.languages import CODE_TO_LANGUAGE
from surya.input.langs import replace_lang_with_code
from surya.schema import OCRResult, TextDetectionResult, LayoutResult, OrderResult
from surya.settings import settings
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL

# Define the function to export OCR results to a .doc file
def export_ocr_to_doc(ocr_result, img_size, doc_path):

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False

    img_width, img_height = img_size

    for line in ocr_result.text_lines:
        p = table.add_row().cells[0].paragraphs[0]
        run = p.add_run(line.text)
        run.font.size = Pt(10)

        bbox = line.bbox
        left = (bbox[0] / img_width) * 100
        top = (bbox[1] / img_height) * 100
        width = ((bbox[2] - bbox[0]) / img_width) * 100
        height = ((bbox[3] - bbox[1]) / img_height) * 100

        cell = table.add_row().cells[0]
        cell.width = Inches(width / 25.4)
        cell.height = Inches(height / 25.4)
        tcPr = cell._element.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    try:
        doc.save(doc_path)
        print(f"Document saved to {doc_path}")
    except PermissionError as e:
        print(f"Permission error: {e}. Try running with administrative privileges or check file permissions.")
    except Exception as e:
        print(f"An error occurred while saving the document: {e}")

# Argument parsing
parser = argparse.ArgumentParser(description="Chạy OCR trên hình ảnh hoặc PDF.")
parser.add_argument("--math", action="store_true", help="Sử dụng mô hình toán học để phát hiện", default=False)

try:
    args = parser.parse_args()
except SystemExit as e:
    print(f"Lỗi phân tích cú pháp đối số: {e}")
    os._exit(e.code)

@st.cache_resource()
def load_det_cached():
    checkpoint = settings.DETECTOR_MATH_MODEL_CHECKPOINT if args.math else settings.DETECTOR_MODEL_CHECKPOINT
    return load_model(checkpoint=checkpoint), load_processor(checkpoint=checkpoint)

@st.cache_resource()
def load_rec_cached():
    return load_rec_model(), load_rec_processor()

@st.cache_resource()
def load_layout_cached():
    return load_model(checkpoint=settings.LAYOUT_MODEL_CHECKPOINT), load_processor(checkpoint=settings.LAYOUT_MODEL_CHECKPOINT)

@st.cache_resource()
def load_order_cached():
    return load_order_model(), load_order_processor()

def text_detection(img) -> (Image.Image, TextDetectionResult):
    pred = batch_text_detection([img], det_model, det_processor)[0]
    polygons = [p.polygon for p in pred.bboxes]
    det_img = draw_polys_on_image(polygons, img.copy())
    return det_img, pred

def layout_detection(img) -> (Image.Image, LayoutResult):
    _, det_pred = text_detection(img)
    pred = batch_layout_detection([img], layout_model, layout_processor, [det_pred])[0]
    polygons = [p.polygon for p in pred.bboxes]
    labels = [p.label for p in pred.bboxes]
    layout_img = draw_polys_on_image(polygons, img.copy(), labels=labels)
    return layout_img, pred

def order_detection(img) -> (Image.Image, OrderResult):
    _, layout_pred = layout_detection(img)
    bboxes = [l.bbox for l in layout_pred.bboxes]
    pred = batch_ordering([img], [bboxes], order_model, order_processor)[0]
    polys = [l.polygon for l in pred.bboxes]
    positions = [str(l.position) for l in pred.bboxes]
    order_img = draw_polys_on_image(polys, img.copy(), labels=positions, label_font_size=20)
    return order_img, pred

# Function for OCR
def ocr(img, langs: List[str]) -> (Image.Image, OCRResult):
    replace_lang_with_code(langs)
    img_pred = run_ocr([img], [langs], det_model, det_processor, rec_model, rec_processor)[0]

    bboxes = [l.bbox for l in img_pred.text_lines]
    text = [l.text for l in img_pred.text_lines]
    rec_img = draw_text_on_image(bboxes, text, img.size, langs, has_math="_math" in langs)
    return rec_img, img_pred

def open_pdf(pdf_file):
    stream = io.BytesIO(pdf_file.getvalue())
    return pypdfium2.PdfDocument(stream)

@st.cache_data()
def get_page_image(pdf_file, page_num, dpi=96):
    doc = open_pdf(pdf_file)
    renderer = doc.render(
        pypdfium2.PdfBitmap.to_pil,
        page_indices=[page_num - 1],
        scale=dpi / 72,
    )
    png = list(renderer)[0]
    png_image = png.convert("RGB")
    return png_image

@st.cache_data()
def page_count(pdf_file):
    doc = open_pdf(pdf_file)
    return len(doc)

st.set_page_config(layout="wide")
col1, col2 = st.columns([.5, .5])

det_model, det_processor = load_det_cached()
rec_model, rec_processor = load_rec_cached()
layout_model, layout_processor = load_layout_cached()
order_model, order_processor = load_order_cached()

st.markdown("""# Surya OCR""")

in_file = st.sidebar.file_uploader("PDF hoặc hình ảnh:", type=["pdf", "png", "jpg", "jpeg", "gif", "webp"])
languages = st.sidebar.multiselect("Ngôn ngữ", sorted(list(CODE_TO_LANGUAGE.values())), default=["Vietnamese"], max_selections=4)

if in_file is None:
    st.stop()

filetype = in_file.type
whole_image = False
if "pdf" in filetype:
    num_pages = page_count(in_file)
    page_number = st.sidebar.number_input(f"Số trang {num_pages}:", min_value=1, value=1, max_value=num_pages)
    pil_image = get_page_image(in_file, page_number)
else:
    pil_image = Image.open(in_file).convert("RGB")

# Run OCR directly and export to DOC
rec_img, pred = ocr(pil_image, languages)
with col1:
    st.image(rec_img, caption="Kết quả OCR", use_column_width=True)
    json_tab, text_tab = st.tabs(["JSON", "Dòng văn bản (for debugging)"])
    with json_tab:
        st.json(pred.model_dump(), expanded=True)
    with text_tab:
        st.text("\n".join([p.text for p in pred.text_lines]))

# Automatically export the OCR result to a DOC file with a unique file name
output_directory = os.path.join(os.getcwd(), "output_docs")
os.makedirs(output_directory, exist_ok=True)
output_doc_path = os.path.join(output_directory, f"output_{int(time.time())}.doc")
export_ocr_to_doc(pred, pil_image.size, output_doc_path)
st.sidebar.success(f"Document exported successfully to {output_doc_path}")

with col2:
    st.image(pil_image, caption="Hình ảnh đã tải lên", use_column_width=True)
